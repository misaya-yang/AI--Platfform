"""
Document Image Extractor

Unified image extraction from various document formats:
- PDF (via PyMuPDF/fitz)
- DOCX (via python-docx)
- HTML (via BeautifulSoup)

Follows Dify 1.11 approach:
- Auto-extract images from documents
- Associate images with nearby text chunks
- Support images up to 2MB for embedding

Usage:
    extractor = DocumentImageExtractor()
    result = await extractor.extract("document.docx", content_bytes)
    for img in result.images:
        if img.is_embeddable:
            # Process image for embedding
            pass
"""

from __future__ import annotations

import asyncio
import base64
import hashlib
import io
import logging
import os
import re
import threading
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import urljoin

logger = logging.getLogger(__name__)


def _get_fitz():
    """Get fitz module with compatibility for old/new PyMuPDF versions."""
    try:
        import pymupdf as fitz

        return fitz
    except ImportError:
        import fitz

        return fitz


# ============================================================
# Constants
# ============================================================

# Supported image formats for multimodal embedding
EMBEDDABLE_IMAGE_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/bmp",
    "image/webp",
    "image/gif",
    "image/svg+xml",
    "image/avif",
}

# Maximum image size for embedding (2MB - Dify standard)
MAX_IMAGE_SIZE_BYTES = 2 * 1024 * 1024

# Maximum document size for processing (100MB)
MAX_DOCUMENT_SIZE_BYTES = 100 * 1024 * 1024

# Minimum image dimensions (skip tiny images)
MIN_IMAGE_WIDTH = 50
MIN_IMAGE_HEIGHT = 50

# Context extraction settings
DEFAULT_CONTEXT_CHARS = 500
MAX_ASSOCIATION_DISTANCE = 500

# Timeout settings
PDF_PAGE_TIMEOUT_SECONDS = 60
HTTP_DOWNLOAD_TIMEOUT_SECONDS = 10.0

# Parallel processing settings
MAX_PARALLEL_WORKERS = 4
MIN_PAGES_FOR_PARALLEL = 3

# Extension to MIME type mapping
EXTENSION_TO_MIME = {
    "png": "image/png",
    "jpeg": "image/jpeg",
    "jpg": "image/jpeg",
    "jpe": "image/jpeg",
    "bmp": "image/bmp",
    "webp": "image/webp",
    "gif": "image/gif",
    "tiff": "image/tiff",
    "tif": "image/tiff",
    "svg": "image/svg+xml",
    "avif": "image/avif",
}


# ============================================================
# Data Classes
# ============================================================


@dataclass
class ExtractedImage:
    """Represents an image extracted from a document."""

    image_id: str  # Unique ID (SHA256 hash of content, 16 chars)
    content: bytes  # Raw image bytes
    mime_type: str  # MIME type (image/png, etc.)
    width: int  # Image width in pixels
    height: int  # Image height in pixels
    source_location: str  # Where found (page number, section, etc.)
    context_text: str = ""  # Surrounding text context
    alt_text: str = ""  # Alt text if available
    filename: str = ""  # Original filename if known
    char_offset: int = -1  # Character offset in document (-1 if unknown)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def size_bytes(self) -> int:
        """Size of image in bytes."""
        return len(self.content)

    @property
    def is_embeddable(self) -> bool:
        """Check if image can be embedded (meets size and dimension requirements)."""
        return (
            self.mime_type in EMBEDDABLE_IMAGE_TYPES
            and self.size_bytes <= MAX_IMAGE_SIZE_BYTES
            and self.width >= MIN_IMAGE_WIDTH
            and self.height >= MIN_IMAGE_HEIGHT
        )

    @property
    def aspect_ratio(self) -> float:
        """Image aspect ratio (width/height). Returns 0.0 if height is 0."""
        if self.height == 0:
            return 0.0
        return self.width / self.height

    def to_dict(self, include_context_length: int = 200) -> dict[str, Any]:
        """Convert to dictionary (without content for serialization).

        Args:
            include_context_length: Maximum length of context_text to include (default: 200)

        Returns:
            Dictionary representation of the image metadata
        """
        return {
            "image_id": self.image_id,
            "mime_type": self.mime_type,
            "width": self.width,
            "height": self.height,
            "source_location": self.source_location,
            "size_bytes": self.size_bytes,
            "context_text": self.context_text[:include_context_length] if self.context_text else "",
            "alt_text": self.alt_text,
            "filename": self.filename,
            "char_offset": self.char_offset,
            "is_embeddable": self.is_embeddable,
            "aspect_ratio": round(self.aspect_ratio, 2),
            "metadata": self.metadata,
        }


@dataclass
class DocumentExtractionResult:
    """Result of document extraction containing text and images."""

    text: str  # Extracted text content
    images: list[ExtractedImage]  # Extracted images
    document_type: str  # pdf, docx, html, etc.
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def embeddable_images(self) -> list[ExtractedImage]:
        """Get only images that can be embedded."""
        return [img for img in self.images if img.is_embeddable]

    @property
    def total_images(self) -> int:
        return len(self.images)

    @property
    def embeddable_image_count(self) -> int:
        return len(self.embeddable_images)

    def associate_images_with_text(
        self,
        chunks: list[dict[str, Any]],
        max_distance: int = MAX_ASSOCIATION_DISTANCE,
    ) -> dict[str, list[ExtractedImage]]:
        """Associate images with text chunks based on proximity.

        Args:
            chunks: List of text chunks with 'start' and 'end' char offsets
            max_distance: Maximum character distance for association (default: 500)

        Returns:
            Dict mapping chunk_id to list of associated images
        """
        associations: dict[str, list[ExtractedImage]] = {}

        for img in self.embeddable_images:
            if img.char_offset < 0:
                continue

            for chunk in chunks:
                chunk_id = chunk.get("chunk_id", chunk.get("id", ""))
                chunk_start = chunk.get("start", 0)
                chunk_end = chunk.get("end", len(self.text))

                # Check if image is within or near this chunk
                if chunk_start - max_distance <= img.char_offset <= chunk_end + max_distance:
                    if chunk_id not in associations:
                        associations[chunk_id] = []
                    associations[chunk_id].append(img)

        return associations


# ============================================================
# Image Utility Functions
# ============================================================


def generate_image_id(content: bytes) -> str:
    """Generate unique ID for image based on content hash.

    Args:
        content: Raw image bytes

    Returns:
        16-character hex string (first 16 chars of SHA256 hash)
    """
    return hashlib.sha256(content).hexdigest()[:16]


def _is_valid_image_dimensions(width: int, height: int, min_width: int, min_height: int) -> bool:
    """Check if image dimensions meet minimum requirements.

    Args:
        width: Image width in pixels
        height: Image height in pixels
        min_width: Minimum acceptable width
        min_height: Minimum acceptable height

    Returns:
        True if dimensions are valid, False otherwise
    """
    return width >= min_width and height >= min_height


def detect_mime_type(content: bytes) -> str:
    """Detect image MIME type from magic bytes.

    Args:
        content: Raw file bytes

    Returns:
        MIME type string (e.g., 'image/png') or 'application/octet-stream' if unknown
    """
    if len(content) < 4:
        return "application/octet-stream"

    # PNG
    if len(content) >= 8 and content[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"

    # JPEG
    if content[:2] == b"\xff\xd8":
        return "image/jpeg"

    # GIF
    if len(content) >= 6 and content[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"

    # BMP
    if content[:2] == b"BM":
        return "image/bmp"

    # WebP
    if len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "image/webp"

    # TIFF
    if content[:4] in (b"II*\x00", b"MM\x00*"):
        return "image/tiff"

    # SVG (XML-based)
    if b"<svg" in content[:1024].lower() or b"<?xml" in content[:100]:
        return "image/svg+xml"

    # AVIF
    if len(content) >= 12 and content[4:8] == b"ftyp":
        # Check for avif/avis brand
        if b"avif" in content[8:16] or b"avis" in content[8:16]:
            return "image/avif"

    # ICO
    if len(content) >= 4 and content[:4] == b"\x00\x00\x01\x00":
        return "image/x-icon"

    return "application/octet-stream"


def get_image_dimensions(content: bytes, mime_type: str) -> tuple[int, int]:
    """Get image dimensions without loading full image.

    Args:
        content: Raw image bytes
        mime_type: MIME type of the image

    Returns:
        Tuple of (width, height) in pixels. Returns (0, 0) if dimensions cannot be determined.
    """
    try:
        from PIL import Image

        with Image.open(io.BytesIO(content)) as img:
            return img.size
    except Exception as e:
        logger.debug(f"PIL failed to read image dimensions: {e}, falling back to header parsing")
        # Fallback: parse header manually for common formats
        return _parse_dimensions_from_header(content, mime_type)


def _parse_dimensions_from_header(content: bytes, mime_type: str) -> tuple[int, int]:
    """Parse image dimensions from header bytes.

    Args:
        content: Raw image bytes
        mime_type: MIME type of the image

    Returns:
        Tuple of (width, height) in pixels. Returns (0, 0) if parsing fails.
    """
    try:
        if mime_type == "image/png" and len(content) >= 24:
            # PNG: width at bytes 16-20, height at 20-24
            width = int.from_bytes(content[16:20], "big")
            height = int.from_bytes(content[20:24], "big")
            return width, height

        elif mime_type in ("image/jpeg", "image/jpg") and len(content) >= 4:
            # JPEG: Need to find SOF0 marker
            data = io.BytesIO(content)
            data.read(2)  # Skip SOI
            while True:
                marker = data.read(2)
                if len(marker) < 2:
                    break
                if marker[0] != 0xFF:
                    break
                if marker[1] in (0xC0, 0xC1, 0xC2):  # SOF markers
                    data.read(3)  # Skip length and precision
                    height_bytes = data.read(2)
                    width_bytes = data.read(2)
                    if len(height_bytes) == 2 and len(width_bytes) == 2:
                        height = int.from_bytes(height_bytes, "big")
                        width = int.from_bytes(width_bytes, "big")
                        return width, height
                    break
                elif marker[1] == 0xD9:  # EOI
                    break
                else:
                    length_bytes = data.read(2)
                    if len(length_bytes) < 2:
                        break
                    length = int.from_bytes(length_bytes, "big")
                    if length < 2:
                        break
                    data.read(length - 2)
            return 0, 0

        elif mime_type == "image/gif" and len(content) >= 10:
            # GIF: width at bytes 6-8, height at 8-10
            width = int.from_bytes(content[6:8], "little")
            height = int.from_bytes(content[8:10], "little")
            return width, height

        elif mime_type == "image/bmp" and len(content) >= 26:
            # BMP: width at bytes 18-22, height at 22-26
            width = int.from_bytes(content[18:22], "little")
            height = abs(int.from_bytes(content[22:26], "little", signed=True))
            return width, height

    except Exception as e:
        logger.debug(f"Failed to parse dimensions from header for {mime_type}: {e}")

    return 0, 0


# ============================================================
# Document Type Extractors
# ============================================================


class PDFExtractor:
    """Extract images from PDF documents using PyMuPDF."""

    def __init__(
        self,
        min_width: int = MIN_IMAGE_WIDTH,
        min_height: int = MIN_IMAGE_HEIGHT,
        context_chars: int = DEFAULT_CONTEXT_CHARS,
    ):
        """Initialize PDF extractor.

        Args:
            min_width: Minimum image width to extract (default: 50)
            min_height: Minimum image height to extract (default: 50)
            context_chars: Number of context characters to extract (default: 500)

        Raises:
            ValueError: If dimensions are invalid
        """
        if min_width <= 0 or min_height <= 0:
            raise ValueError("Image dimensions must be positive")
        if context_chars < 0:
            raise ValueError("Context chars must be non-negative")

        self.min_width = min_width
        self.min_height = min_height
        self.context_chars = context_chars

    async def extract(self, content: bytes) -> DocumentExtractionResult:
        """Extract text and images from PDF with parallel processing."""
        try:
            _get_fitz()
        except ImportError:
            logger.warning("PyMuPDF not installed, using fallback")
            return await self._fallback_extract(content)

        return await asyncio.to_thread(self._extract_sync_parallel, content)

    def _extract_page_images_parallel(
        self,
        content: bytes,
        page_num: int,
        page_text: str,
        char_offset: int,
        seen_hashes: set,
        seen_lock: threading.Lock,
    ) -> list[ExtractedImage]:
        """Extract images from a single page using an isolated document instance.

        Note: This method opens a separate document instance per thread to avoid
        threading issues with PyMuPDF. For large PDFs, this increases memory usage.

        Args:
            content: Full PDF content bytes
            page_num: Page number to extract (0-indexed)
            page_text: Text content of this page
            char_offset: Character offset of this page in full document
            seen_hashes: Set of seen image hashes (shared, thread-safe)
            seen_lock: Lock for seen_hashes access

        Returns:
            List of extracted images from this page
        """
        fitz = _get_fitz()

        page_images: list[ExtractedImage] = []
        doc = None
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            page = doc[page_num]

            for img_index, img in enumerate(page.get_images(full=True)):
                try:
                    xref = img[0]
                    base_img = doc.extract_image(xref)

                    if not base_img:
                        continue

                    img_bytes = base_img["image"]

                    # Compute hash first (outside lock to reduce contention)
                    img_hash = hashlib.sha256(img_bytes).hexdigest()

                    # Deduplicate by hash (thread-safe)
                    with seen_lock:
                        if img_hash in seen_hashes:
                            continue
                        seen_hashes.add(img_hash)

                    img_ext = base_img.get("ext", "png")
                    mime_type = EXTENSION_TO_MIME.get(img_ext, "image/png")

                    # Get dimensions
                    width = base_img.get("width", 0)
                    height = base_img.get("height", 0)

                    # Skip small images
                    if not _is_valid_image_dimensions(
                        width, height, self.min_width, self.min_height
                    ):
                        logger.debug(
                            f"Skipping small image on page {page_num + 1}: {width}x{height}"
                        )
                        continue

                    # Get context text
                    context = page_text[: self.context_chars] if page_text else ""

                    page_images.append(
                        ExtractedImage(
                            image_id=generate_image_id(img_bytes),
                            content=img_bytes,
                            mime_type=mime_type,
                            width=width,
                            height=height,
                            source_location=f"page_{page_num + 1}",
                            context_text=context,
                            char_offset=char_offset,
                            metadata={
                                "page_number": page_num + 1,
                                "image_index": img_index,
                                "xref": xref,
                            },
                        )
                    )

                except Exception as e:
                    logger.warning(
                        f"Failed to extract image {img_index} from page {page_num + 1}: {e}"
                    )

            return page_images
        except Exception as e:
            logger.error(f"Failed to process page {page_num + 1}: {e}")
            return []
        finally:
            if doc:
                doc.close()

    def _extract_sync_parallel(self, content: bytes) -> DocumentExtractionResult:
        """Synchronous extraction with parallel image processing.

        Args:
            content: PDF file bytes

        Returns:
            DocumentExtractionResult containing text and images
        """
        fitz = _get_fitz()
        from concurrent.futures import ThreadPoolExecutor, TimeoutError, as_completed

        doc = None
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            total_pages = len(doc)
            text_parts: list[str] = []

            # Extract all page text sequentially for stable offsets
            for page_num in range(total_pages):
                page = doc[page_num]
                text_parts.append(page.get_text("text"))

            # Compute character offsets (account for "\n\n" join)
            char_offsets: list[int] = []
            running = 0
            for idx, page_text in enumerate(text_parts):
                char_offsets.append(running)
                running += len(page_text)
                if idx < total_pages - 1:
                    running += 2  # "\n\n" separator

            all_images: list[ExtractedImage] = []
            seen_hashes: set = set()
            seen_lock = threading.Lock()
            extraction_errors = 0

            # For small documents, use sequential processing
            if total_pages < MIN_PAGES_FOR_PARALLEL:
                logger.debug(f"Using sequential processing for {total_pages} pages")
                for page_num in range(total_pages):
                    page_images = self._extract_page_images_parallel(
                        content,
                        page_num,
                        text_parts[page_num],
                        char_offsets[page_num],
                        seen_hashes,
                        seen_lock,
                    )
                    all_images.extend(page_images)
            else:
                # Parallel processing for larger documents (isolated document per worker)
                max_workers = min(MAX_PARALLEL_WORKERS, (os.cpu_count() or 2))
                logger.info(
                    f"Extracting images from {total_pages} pages with {max_workers} workers"
                )

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    future_to_page = {
                        executor.submit(
                            self._extract_page_images_parallel,
                            content,
                            page_num,
                            text_parts[page_num],
                            char_offsets[page_num],
                            seen_hashes,
                            seen_lock,
                        ): page_num
                        for page_num in range(total_pages)
                    }

                    for future in as_completed(future_to_page):
                        page_num = future_to_page[future]
                        try:
                            page_images = future.result(timeout=PDF_PAGE_TIMEOUT_SECONDS)
                            all_images.extend(page_images)
                        except TimeoutError:
                            extraction_errors += 1
                            logger.error(
                                f"Timeout extracting images from page {page_num + 1} after {PDF_PAGE_TIMEOUT_SECONDS}s"
                            )
                        except Exception as e:
                            extraction_errors += 1
                            logger.error(f"Failed to extract images from page {page_num + 1}: {e}")

            return DocumentExtractionResult(
                text="\n\n".join(text_parts),
                images=all_images,
                document_type="pdf",
                metadata={
                    "page_count": total_pages,
                    "total_images_found": len(all_images),
                    "extraction_errors": extraction_errors,
                },
            )
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="pdf",
                metadata={"error": str(e)},
            )
        finally:
            if doc:
                doc.close()

    def _extract_sync(self, content: bytes) -> DocumentExtractionResult:
        """Synchronous extraction (runs in thread) - legacy method for backward compatibility."""
        return self._extract_sync_parallel(content)

    async def _fallback_extract(self, content: bytes) -> DocumentExtractionResult:
        """Fallback using pypdf if PyMuPDF not available."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="pdf",
                metadata={"error": "No PDF library available"},
            )

        reader = PdfReader(io.BytesIO(content))
        text_parts = [page.extract_text() or "" for page in reader.pages]

        return DocumentExtractionResult(
            text="\n\n".join(text_parts),
            images=[],  # pypdf doesn't easily extract images
            document_type="pdf",
            metadata={"page_count": len(reader.pages)},
        )


class DOCXExtractor:
    """Extract images from DOCX documents using python-docx."""

    def __init__(
        self,
        min_width: int = MIN_IMAGE_WIDTH,
        min_height: int = MIN_IMAGE_HEIGHT,
        context_chars: int = DEFAULT_CONTEXT_CHARS,
    ):
        """Initialize DOCX extractor.

        Args:
            min_width: Minimum image width to extract (default: 50)
            min_height: Minimum image height to extract (default: 50)
            context_chars: Number of context characters to extract (default: 500)

        Raises:
            ValueError: If dimensions are invalid
        """
        if min_width <= 0 or min_height <= 0:
            raise ValueError("Image dimensions must be positive")
        if context_chars < 0:
            raise ValueError("Context chars must be non-negative")

        self.min_width = min_width
        self.min_height = min_height
        self.context_chars = context_chars

    async def extract(self, content: bytes) -> DocumentExtractionResult:
        """Extract text and images from DOCX."""
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="docx",
                metadata={"error": "python-docx not installed"},
            )

        return await asyncio.to_thread(self._extract_sync, content)

    def _extract_sync(self, content: bytes) -> DocumentExtractionResult:
        """Synchronous extraction.

        Args:
            content: DOCX file bytes

        Returns:
            DocumentExtractionResult containing text and images
        """
        from docx import Document

        extraction_errors = 0
        try:
            doc = Document(io.BytesIO(content))
            text_parts: list[str] = []
            images: list[ExtractedImage] = []
            seen_hashes: set = set()

            # Extract text from paragraphs
            char_offset = 0
            for para in doc.paragraphs:
                text_parts.append(para.text)
                char_offset += len(para.text) + 1

            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        img_part = rel.target_part
                        img_bytes = img_part.blob

                        # Deduplicate by hash
                        img_hash = hashlib.sha256(img_bytes).hexdigest()
                        if img_hash in seen_hashes:
                            continue
                        seen_hashes.add(img_hash)

                        # Detect type and dimensions
                        mime_type = detect_mime_type(img_bytes)
                        width, height = get_image_dimensions(img_bytes, mime_type)

                        # Skip small images
                        if not _is_valid_image_dimensions(
                            width, height, self.min_width, self.min_height
                        ):
                            logger.debug(f"Skipping small DOCX image: {width}x{height}")
                            continue

                        # Get filename from content type
                        filename = getattr(img_part, "partname", "").split("/")[-1]

                        images.append(
                            ExtractedImage(
                                image_id=generate_image_id(img_bytes),
                                content=img_bytes,
                                mime_type=mime_type,
                                width=width,
                                height=height,
                                source_location="document_body",
                                filename=filename,
                                metadata={"relationship_id": rel.rId},
                            )
                        )

                    except Exception as e:
                        extraction_errors += 1
                        logger.warning(f"Failed to extract DOCX image: {e}")

            full_text = "\n".join(text_parts)

            return DocumentExtractionResult(
                text=full_text,
                images=images,
                document_type="docx",
                metadata={
                    "paragraph_count": len(doc.paragraphs),
                    "total_images_found": len(images),
                    "extraction_errors": extraction_errors,
                },
            )
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="docx",
                metadata={"error": str(e)},
            )


class HTMLExtractor:
    """Extract images from HTML documents using BeautifulSoup."""

    def __init__(
        self,
        min_width: int = MIN_IMAGE_WIDTH,
        min_height: int = MIN_IMAGE_HEIGHT,
        context_chars: int = DEFAULT_CONTEXT_CHARS,
        download_images: bool = False,
        base_url: str | None = None,
    ):
        """Initialize HTML extractor.

        Args:
            min_width: Minimum image width to extract (default: 50)
            min_height: Minimum image height to extract (default: 50)
            context_chars: Number of context characters to extract (default: 500)
            download_images: Whether to download external images (default: False)
            base_url: Base URL for resolving relative image URLs (optional)

        Raises:
            ValueError: If dimensions are invalid
        """
        if min_width <= 0 or min_height <= 0:
            raise ValueError("Image dimensions must be positive")
        if context_chars < 0:
            raise ValueError("Context chars must be non-negative")

        self.min_width = min_width
        self.min_height = min_height
        self.context_chars = context_chars
        self.download_images = download_images
        self.base_url = base_url

    async def extract(
        self,
        content: bytes | str,
        base_url: str | None = None,
    ) -> DocumentExtractionResult:
        """Extract text and images from HTML.

        Args:
            content: HTML content as bytes or string
            base_url: Base URL for resolving relative image URLs (optional)

        Returns:
            DocumentExtractionResult containing text and images
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("BeautifulSoup not installed, cannot extract HTML")
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="html",
                metadata={"error": "BeautifulSoup not installed"},
            )

        extraction_errors = 0
        try:
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="ignore")

            soup = BeautifulSoup(content, "html.parser")

            # Extract text
            for script in soup(["script", "style", "noscript"]):
                script.decompose()
            text = soup.get_text(separator="\n", strip=True)

            # Extract images
            images: list[ExtractedImage] = []
            seen_hashes: set = set()

            resolved_base_url = base_url or self.base_url
            total_img_tags = len(soup.find_all("img"))

            for img_tag in soup.find_all("img"):
                try:
                    src = img_tag.get("src", "")
                    alt = img_tag.get("alt", "")

                    if not src:
                        continue

                    # Handle data URIs
                    if src.startswith("data:"):
                        img_data = self._parse_data_uri(src)
                        if img_data:
                            content_bytes, mime_type = img_data

                            # Deduplicate
                            img_hash = hashlib.sha256(content_bytes).hexdigest()
                            if img_hash in seen_hashes:
                                continue
                            seen_hashes.add(img_hash)

                            width, height = get_image_dimensions(content_bytes, mime_type)

                            if _is_valid_image_dimensions(
                                width, height, self.min_width, self.min_height
                            ):
                                # Get surrounding text
                                context = self._get_surrounding_text(img_tag, self.context_chars)

                                images.append(
                                    ExtractedImage(
                                        image_id=generate_image_id(content_bytes),
                                        content=content_bytes,
                                        mime_type=mime_type,
                                        width=width,
                                        height=height,
                                        source_location="inline_data_uri",
                                        alt_text=alt,
                                        context_text=context,
                                    )
                                )
                            else:
                                logger.debug(
                                    f"Skipping small HTML data URI image: {width}x{height}"
                                )

                    elif self.download_images:
                        # Download external images (async)
                        img_result = await self._download_image(src, resolved_base_url)
                        if img_result:
                            content_bytes, mime_type = img_result

                            img_hash = hashlib.sha256(content_bytes).hexdigest()
                            if img_hash in seen_hashes:
                                continue
                            seen_hashes.add(img_hash)

                            width, height = get_image_dimensions(content_bytes, mime_type)

                            if _is_valid_image_dimensions(
                                width, height, self.min_width, self.min_height
                            ):
                                context = self._get_surrounding_text(img_tag, self.context_chars)
                                filename = src.split("/")[-1].split("?")[0]

                                images.append(
                                    ExtractedImage(
                                        image_id=generate_image_id(content_bytes),
                                        content=content_bytes,
                                        mime_type=mime_type,
                                        width=width,
                                        height=height,
                                        source_location=src,
                                        alt_text=alt,
                                        filename=filename,
                                        context_text=context,
                                    )
                                )
                            else:
                                logger.debug(
                                    f"Skipping small HTML downloaded image: {width}x{height}"
                                )

                except Exception as e:
                    extraction_errors += 1
                    logger.warning(f"Failed to extract HTML image from {src[:100]}: {e}")

            return DocumentExtractionResult(
                text=text,
                images=images,
                document_type="html",
                metadata={
                    "total_img_tags": total_img_tags,
                    "extracted_images": len(images),
                    "extraction_errors": extraction_errors,
                },
            )
        except Exception as e:
            logger.error(f"Failed to extract HTML content: {e}")
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type="html",
                metadata={"error": str(e)},
            )

    def _parse_data_uri(self, data_uri: str) -> tuple[bytes, str] | None:
        """Parse a data URI into bytes and MIME type.

        Args:
            data_uri: Data URI string (e.g., 'data:image/png;base64,...')

        Returns:
            Tuple of (content_bytes, mime_type) or None if parsing fails
        """
        match = re.match(r"data:([^;,]+)?(?:;base64)?,(.+)", data_uri)
        if not match:
            logger.debug("Invalid data URI format")
            return None

        mime_type = match.group(1) or "application/octet-stream"
        data = match.group(2)

        try:
            content = base64.b64decode(data)
            return content, mime_type
        except Exception as e:
            logger.debug(f"Failed to decode base64 data URI: {e}")
            return None

    def _get_surrounding_text(self, element, max_chars: int) -> str:
        """Get text surrounding an element."""
        parts = []

        # Get preceding text
        prev = element.find_previous(string=True)
        if prev:
            parts.append(str(prev).strip()[: max_chars // 2])

        # Get following text
        next_elem = element.find_next(string=True)
        if next_elem:
            parts.append(str(next_elem).strip()[: max_chars // 2])

        return " ".join(parts)

    async def _download_image(
        self, url: str, base_url: str | None = None
    ) -> tuple[bytes, str] | None:
        """Download an image from URL.

        Args:
            url: Image URL (absolute or relative)
            base_url: Base URL for resolving relative URLs (optional)

        Returns:
            Tuple of (content_bytes, mime_type) or None if download fails
        """
        try:
            import httpx

            # Resolve relative URLs
            resolved_base_url = base_url or self.base_url
            if resolved_base_url and not url.startswith(("http://", "https://")):
                url = urljoin(resolved_base_url, url)

            if not url.startswith(("http://", "https://")):
                logger.debug(f"Skipping non-HTTP URL: {url[:100]}")
                return None

            # Add headers to avoid being blocked as bot
            headers = {
                "User-Agent": "Mozilla/5.0 (compatible; AI-Gateway/1.0; +https://github.com/ai-gateway)",
                "Accept": "image/avif,image/webp,image/png,image/jpeg,*/*",
            }

            async with httpx.AsyncClient(
                timeout=HTTP_DOWNLOAD_TIMEOUT_SECONDS,
                follow_redirects=True,
                headers=headers,
            ) as client:
                response = await client.get(url)

                if response.status_code == 200:
                    content = response.content

                    # Check content size
                    if len(content) > MAX_IMAGE_SIZE_BYTES:
                        logger.info(
                            f"Downloaded image too large: {len(content)} bytes from {url[:100]}"
                        )
                        return None

                    mime_type = response.headers.get("content-type", "").split(";")[0]
                    if not mime_type.startswith("image/"):
                        mime_type = detect_mime_type(content)

                    return content, mime_type
                else:
                    logger.debug(
                        f"Failed to download image: HTTP {response.status_code} from {url[:100]}"
                    )
                    return None

        except Exception as e:
            logger.warning(f"Failed to download image from {url[:100]}: {e}")
            return None


# ============================================================
# Unified Document Image Extractor
# ============================================================


class DocumentImageExtractor:
    """Unified image extractor supporting multiple document formats.

    Supported formats:
    - PDF (.pdf) - via PyMuPDF
    - DOCX (.docx) - via python-docx
    - HTML (.html, .htm) - via BeautifulSoup

    Usage:
        extractor = DocumentImageExtractor()

        # Extract from bytes
        result = await extractor.extract("document.pdf", pdf_bytes)

        # Get embeddable images
        for img in result.embeddable_images:
            print(f"Image: {img.image_id}, Size: {img.size_bytes}")

        # Associate images with text chunks
        associations = result.associate_images_with_text(chunks)
    """

    def __init__(
        self,
        min_width: int = MIN_IMAGE_WIDTH,
        min_height: int = MIN_IMAGE_HEIGHT,
        context_chars: int = DEFAULT_CONTEXT_CHARS,
        download_html_images: bool = False,
    ):
        """Initialize the extractor.

        Args:
            min_width: Minimum image width to extract (default: 50)
            min_height: Minimum image height to extract (default: 50)
            context_chars: Characters of context to extract (default: 500)
            download_html_images: Whether to download external images in HTML (default: False)

        Raises:
            ValueError: If dimensions are invalid
        """
        if min_width <= 0 or min_height <= 0:
            raise ValueError("Image dimensions must be positive")
        if context_chars < 0:
            raise ValueError("Context chars must be non-negative")

        self.min_width = min_width
        self.min_height = min_height
        self.context_chars = context_chars
        self.download_html_images = download_html_images

        # Initialize specialized extractors
        self._pdf_extractor = PDFExtractor(min_width, min_height, context_chars)
        self._docx_extractor = DOCXExtractor(min_width, min_height, context_chars)
        self._html_extractor = HTMLExtractor(
            min_width,
            min_height,
            context_chars,
            download_images=download_html_images,
        )

    def _detect_document_type(self, filename: str, content: bytes) -> str:
        """Detect document type from filename or content.

        Args:
            filename: Document filename
            content: Document content bytes

        Returns:
            Document type: 'pdf', 'docx', 'html', or 'unknown'
        """
        ext = Path(filename).suffix.lower() if filename else ""

        # Check by extension first (fast path)
        if ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        elif ext in (".html", ".htm"):
            return "html"

        # Fallback to content detection (with boundary checks)
        if len(content) >= 4:
            # PDF magic bytes
            if content[:4] == b"%PDF":
                return "pdf"
            # ZIP/DOCX magic bytes (DOCX is a ZIP file)
            elif content[:4] == b"PK\x03\x04":
                # Could be DOCX or other ZIP files
                # DOCX typically contains word/ directory
                if b"word/" in content[:2000]:
                    return "docx"
                # Generic ZIP - treat as unknown
                return "unknown"

        # HTML detection (check first 1KB safely)
        try:
            check_length = min(len(content), 1000)
            if (
                b"<html" in content[:check_length].lower()
                or b"<!doctype html" in content[:check_length].lower()
            ):
                return "html"
        except Exception as e:
            logger.debug(f"Failed HTML detection: {e}")

        return "unknown"

    async def extract(
        self,
        filename: str,
        content: bytes,
        document_type: str | None = None,
        base_url: str | None = None,
    ) -> DocumentExtractionResult:
        """Extract text and images from a document.

        Args:
            filename: Document filename (used for type detection)
            content: Document content bytes
            document_type: Override auto-detected type (pdf, docx, html)
            base_url: Base URL for resolving relative image URLs (HTML only)

        Returns:
            DocumentExtractionResult with text and images

        Raises:
            ValueError: If content is empty or too large
        """
        # Input validation
        if not content:
            raise ValueError("Document content cannot be empty")

        if len(content) > MAX_DOCUMENT_SIZE_BYTES:
            raise ValueError(
                f"Document too large: {len(content)} bytes (max: {MAX_DOCUMENT_SIZE_BYTES})"
            )

        if not filename:
            logger.warning("No filename provided, relying on content detection")

        doc_type = document_type or self._detect_document_type(filename, content)
        logger.info(f"Extracting {doc_type} document: {filename} ({len(content)} bytes)")

        try:
            if doc_type == "pdf":
                return await self._pdf_extractor.extract(content)

            elif doc_type == "docx":
                return await self._docx_extractor.extract(content)

            elif doc_type == "html":
                return await self._html_extractor.extract(content, base_url=base_url)

            else:
                logger.warning(f"Unsupported document type: {doc_type} for {filename}")
                return DocumentExtractionResult(
                    text="",
                    images=[],
                    document_type=doc_type,
                    metadata={"error": f"Unsupported document type: {doc_type}"},
                )
        except Exception as e:
            logger.error(f"Failed to extract document {filename}: {e}")
            return DocumentExtractionResult(
                text="",
                images=[],
                document_type=doc_type,
                metadata={"error": str(e)},
            )

    async def extract_from_file(
        self,
        file_path: str | Path,
        base_url: str | None = None,
    ) -> DocumentExtractionResult:
        """Extract from a file path.

        Args:
            file_path: Path to the document file
            base_url: Base URL for HTML images

        Returns:
            DocumentExtractionResult

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is too large or empty
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.is_file():
            raise ValueError(f"Not a file: {file_path}")

        file_size = path.stat().st_size
        if file_size == 0:
            raise ValueError(f"File is empty: {file_path}")

        if file_size > MAX_DOCUMENT_SIZE_BYTES:
            raise ValueError(f"File too large: {file_size} bytes (max: {MAX_DOCUMENT_SIZE_BYTES})")

        content = path.read_bytes()
        return await self.extract(path.name, content, base_url=base_url)


# ============================================================
# Convenience Functions
# ============================================================


async def extract_images_from_document(
    filename: str,
    content: bytes,
    min_size: int = MIN_IMAGE_WIDTH,
) -> list[ExtractedImage]:
    """Convenience function to extract embeddable images from a document.

    Args:
        filename: Document filename
        content: Document content bytes
        min_size: Minimum image dimension (default: 50)

    Returns:
        List of embeddable images

    Raises:
        ValueError: If content is empty or invalid parameters
    """
    if min_size <= 0:
        raise ValueError("min_size must be positive")

    extractor = DocumentImageExtractor(min_width=min_size, min_height=min_size)
    result = await extractor.extract(filename, content)
    return result.embeddable_images
