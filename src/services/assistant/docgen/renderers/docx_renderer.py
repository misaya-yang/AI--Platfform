"""DocxRenderer — pure-python (python-docx).

Three paths the plan lists are:

1. ``docx-js`` via Node sidecar — **preferred from-scratch path**.
2. ``python-docx`` — compatible fallback; this module's main path.
3. OOXML unpack + edit + pack — used for editing existing .docx.

For the main assistant backend we don't run Node at first, so path 2 is
what we implement now. The docx-js sidecar lives in
``scripts/node_renderers/docx_render.js`` and is wired through the
sandbox client in Phase 2.
"""

from __future__ import annotations

import re
import time
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

from ..ir import (
    BulletBlock,
    ChartBlock,
    CodeBlock,
    DocxIR,
    HeadingBlock,
    ImageBlock,
    ParagraphBlock,
    QuoteBlock,
    TableBlock,
)
from .base import BaseRenderer, RenderError, RenderResult


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _hex_to_rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


class DocxRenderer(BaseRenderer):
    format = "docx"

    def _apply_theme(self, doc: Document, ir: DocxIR) -> None:
        try:
            style = doc.styles["Normal"]
            style.font.name = ir.theme.font_primary.family
            style.font.size = Pt(ir.theme.font_primary.size_pt)
        except Exception:
            pass

    def _render_paragraph(self, doc: Document, block: ParagraphBlock) -> None:
        p = doc.add_paragraph(block.text)
        p.alignment = _ALIGN.get(block.align, WD_ALIGN_PARAGRAPH.LEFT)
        if block.font:
            for run in p.runs:
                run.font.size = Pt(block.font.size_pt)
                run.font.bold = block.font.bold
                run.font.italic = block.font.italic

    def _render_heading(self, doc: Document, block: HeadingBlock) -> None:
        h = doc.add_heading(block.text, level=block.level)
        if block.align == "center":
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_bullets(self, doc: Document, block: BulletBlock) -> None:
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            doc.add_paragraph(item, style=style)

    def _render_quote(self, doc: Document, block: QuoteBlock) -> None:
        try:
            p = doc.add_paragraph(block.text, style="Intense Quote")
        except KeyError:
            p = doc.add_paragraph(block.text)
            for run in p.runs:
                run.italic = True
        if block.author:
            doc.add_paragraph(f"— {block.author}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _render_code(self, doc: Document, block: CodeBlock) -> None:
        # python-docx has no first-class code block; approximate with
        # monospaced font + pale grey fill via runs.
        p = doc.add_paragraph()
        run = p.add_run(block.code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _render_table(self, doc: Document, block: TableBlock) -> None:
        if not block.rows:
            return
        n_cols = max(len(r.cells) for r in block.rows)
        table = doc.add_table(rows=len(block.rows), cols=n_cols)
        table.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(block.rows):
            for c_idx in range(n_cols):
                cell = table.rows[r_idx].cells[c_idx]
                if c_idx < len(row.cells):
                    src = row.cells[c_idx]
                    cell.text = src.text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = src.bold or row.is_header
        if block.caption:
            cap = doc.add_paragraph(block.caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True

    def _render_image(self, doc: Document, block: ImageBlock, out_dir: Path) -> None:
        if not block.source_path:
            doc.add_paragraph(f"[image: {block.alt_text} (missing path)]")
            return
        src = Path(block.source_path)
        if not src.is_absolute():
            src = out_dir / src
        if not src.exists():
            doc.add_paragraph(f"[image: {block.alt_text} (not found: {src})]")
            return
        width = Cm(block.width_pt / 28.35) if block.width_pt else Cm(14)
        doc.add_picture(str(src), width=width)
        # Alt text: python-docx needs to patch the XML directly to set descr=
        try:
            self._set_image_alt(doc, block.alt_text)
        except Exception:
            pass

    def _set_image_alt(self, doc: Document, alt: str) -> None:
        # Iterate the last inline picture and set descr/title on the <wp:docPr>.
        from docx.oxml.ns import qn

        body = doc.element.body
        docprs = body.iter(qn("wp:docPr"))
        last = None
        for el in docprs:
            last = el
        if last is not None:
            last.set("descr", alt[:400])
            if not last.get("title"):
                last.set("title", alt[:64])

    def _render_chart_placeholder(self, doc: Document, block: ChartBlock) -> None:
        # Native charts in python-docx require insert XML; keep a textual
        # placeholder, and let the Node sidecar upgrade this in Phase 2.
        doc.add_paragraph(f"[chart: {block.alt_text}]")

    def _render_block(self, doc: Document, block, out_dir: Path) -> None:
        if isinstance(block, HeadingBlock):
            self._render_heading(doc, block)
        elif isinstance(block, ParagraphBlock):
            self._render_paragraph(doc, block)
        elif isinstance(block, BulletBlock):
            self._render_bullets(doc, block)
        elif isinstance(block, QuoteBlock):
            self._render_quote(doc, block)
        elif isinstance(block, CodeBlock):
            self._render_code(doc, block)
        elif isinstance(block, TableBlock):
            self._render_table(doc, block)
        elif isinstance(block, ImageBlock):
            self._render_image(doc, block, out_dir)
        elif isinstance(block, ChartBlock):
            self._render_chart_placeholder(doc, block)
        else:
            doc.add_paragraph(f"[unsupported block: {type(block).__name__}]")

    async def render(self, ir: DocxIR, out_dir: Path) -> RenderResult:
        if not isinstance(ir, DocxIR):
            raise RenderError(f"DocxRenderer got wrong IR type: {type(ir).__name__}")
        started = time.perf_counter()
        doc = Document()
        self._apply_theme(doc, ir)

        doc.core_properties.title = ir.metadata.title
        if ir.metadata.author:
            doc.core_properties.author = ir.metadata.author

        # Title page-ish
        title_heading = doc.add_heading(ir.metadata.title, level=0)
        if ir.metadata.subtitle:
            sub = doc.add_paragraph(ir.metadata.subtitle)
            sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if ir.content.header:
            section = doc.sections[0]
            section.header.paragraphs[0].text = ir.content.header.text
        if ir.content.footer:
            section = doc.sections[0]
            section.footer.paragraphs[0].text = ir.content.footer.text

        for block in ir.content.blocks:
            self._render_block(doc, block, out_dir)

        out_dir.mkdir(parents=True, exist_ok=True)
        safe_name = re.sub(r"[^A-Za-z0-9_\-\.]", "_", ir.metadata.title)[:80] or "document"
        path = out_dir / f"{safe_name}.docx"
        doc.save(str(path))

        return RenderResult(
            path=path,
            doc_type="docx",
            bytes_size=path.stat().st_size,
            duration_ms=int((time.perf_counter() - started) * 1000),
        )

    async def fix(self, ir: DocxIR, critic_findings, out_dir: Path) -> RenderResult:
        # DOCX fixes are OOXML-level (strip placeholders, fix alt text).
        # For Phase 1 we just re-render; Phase 3 adds targeted patches.
        return await self.render(ir, out_dir)
