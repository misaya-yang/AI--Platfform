"""
Document Generation Tool for Assistant Service

Provides document generation capabilities:
- Markdown to Word (.docx) conversion
- Markdown to PDF conversion (optional, requires weasyprint)
- Plain Markdown output

Uses python-docx for Word document generation.
"""

from __future__ import annotations

import base64
import io
import re
import time

from ai_gateway_core.logging import get_logger
from .tool_registry import (
    ToolCallRequest,
    ToolCallResult,
    ToolCategory,
    ToolDefinition,
    ToolExample,
    ToolExecutor,
    ToolParameter,
    ToolRiskLevel,
    register_tool,
)

logger = get_logger(__name__)


# =============================================================================
# Document Generation Tool Definition
# =============================================================================

DOCUMENT_GENERATION_DEFINITION = ToolDefinition(
    name="generate_document",
    description="Generate a formatted document (Word, PDF, or Markdown) from text content. "
    "IMPORTANT: The content parameter must contain COMPLETE, DETAILED document content in Markdown format. "
    "Before calling this tool, you MUST first write out the full document content in your chat response.",
    parameters=[
        ToolParameter(
            name="title",
            type="string",
            description="Title of the document (will be used as filename and document title).",
            required=True,
        ),
        ToolParameter(
            name="content",
            type="string",
            description="COMPLETE document content in Markdown format. This MUST be the FULL detailed content, "
            "not a skeleton or outline. Include: headers (# ## ###), detailed paragraphs, lists, "
            "examples, and analysis. The content should be comprehensive and well-structured. "
            "Minimum 500+ words for typical documents.",
            required=True,
        ),
        ToolParameter(
            name="format",
            type="string",
            description="Output format: 'docx' for Word, 'pdf' for PDF, 'md' for Markdown.",
            required=False,
            default="docx",
            enum=["docx", "pdf", "md"],
        ),
    ],
    category=ToolCategory.GENERATION,
    risk_level=ToolRiskLevel.LOW,
    when_to_use="Use ONLY AFTER you have written out the complete document content in your chat response. "
    "The user should see the full content before you generate the document file. "
    "Good for creating detailed reports, plans, analyses, or documentation.",
    when_not_to_use="Do NOT use with minimal/skeleton content. Do NOT use for simple responses. "
    "Do NOT call immediately - first show the full content in chat, then generate.",
    examples=[
        ToolExample(
            description="Generate a sales report",
            input={
                "title": "Q4 Sales Report",
                "content": "# Q4 Sales Report\n\n## Summary\n- Total revenue: $1.2M\n- Growth: 15%\n\n## Details\n...",
                "format": "docx",
            },
            expected_output="Returns a Word document with the formatted report",
        ),
        ToolExample(
            description="Generate meeting notes",
            input={
                "title": "Team Meeting Notes 2025-01-15",
                "content": "# Team Meeting Notes\n\n**Date:** January 15, 2025\n\n## Attendees\n- Alice\n- Bob\n\n## Action Items\n1. Review project timeline\n2. Update documentation",
                "format": "docx",
            },
            expected_output="Returns a Word document with meeting notes",
        ),
    ],
    timeout_seconds=30,
)


# =============================================================================
# Markdown to Document Converter
# =============================================================================


class MarkdownToDocxConverter:
    """Convert Markdown to Word document using python-docx."""

    def __init__(self):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt

            self._docx_available = True
        except ImportError:
            self._docx_available = False
            logger.warning("python-docx not installed, Word generation disabled")

    @property
    def is_available(self) -> bool:
        return self._docx_available

    def convert(self, title: str, markdown_content: str) -> bytes:
        """Convert markdown to Word document bytes."""
        if not self._docx_available:
            raise RuntimeError("python-docx is not installed")

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse and add content
        self._parse_markdown(doc, markdown_content)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _parse_markdown(self, doc, content: str):
        """Parse markdown and add to document."""
        from docx.shared import Pt

        lines = content.split("\n")
        in_code_block = False
        code_content = []
        in_list = False
        list_items = []

        for line in lines:
            # Handle code blocks
            if line.startswith("```"):
                if in_code_block:
                    # End code block
                    if code_content:
                        code_para = doc.add_paragraph()
                        code_para.style = "No Spacing"
                        run = code_para.add_run("\n".join(code_content))
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                    code_content = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_content.append(line)
                continue

            # Skip empty lines but end lists
            if not line.strip():
                if in_list and list_items:
                    self._add_list(doc, list_items)
                    list_items = []
                    in_list = False
                continue

            # Handle headers
            if line.startswith("#"):
                if in_list and list_items:
                    self._add_list(doc, list_items)
                    list_items = []
                    in_list = False

                level = len(re.match(r"^#+", line).group())
                text = line.lstrip("#").strip()
                doc.add_heading(text, level=min(level, 9))
                continue

            # Handle bullet lists
            if re.match(r"^[\-\*]\s", line):
                in_list = True
                list_items.append(("bullet", line[2:].strip()))
                continue

            # Handle numbered lists
            if re.match(r"^\d+\.\s", line):
                in_list = True
                text = re.sub(r"^\d+\.\s", "", line).strip()
                list_items.append(("number", text))
                continue

            # Regular paragraph
            if in_list and list_items:
                self._add_list(doc, list_items)
                list_items = []
                in_list = False

            para = doc.add_paragraph()
            self._add_formatted_text(para, line)

        # Handle remaining list items
        if in_list and list_items:
            self._add_list(doc, list_items)

    def _add_list(self, doc, items: list[tuple]):
        """Add a list to the document."""
        for item_type, text in items:
            if item_type == "bullet":
                para = doc.add_paragraph(style="List Bullet")
            else:
                para = doc.add_paragraph(style="List Number")
            self._add_formatted_text(para, text)

    def _add_formatted_text(self, para, text: str):
        """Add text with basic formatting (bold, italic)."""
        # Simple regex-based formatting
        # Handle **bold** and *italic*
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
            else:
                para.add_run(part)


class MarkdownToPdfConverter:
    """Convert Markdown to PDF (requires weasyprint or similar)."""

    def __init__(self):
        self._available = False
        try:
            import markdown

            self._md = markdown.Markdown(extensions=["tables", "fenced_code"])
            # Check for PDF generation capability
            try:
                from weasyprint import HTML

                self._available = True
            except ImportError:
                logger.info("weasyprint not installed, PDF generation disabled")
        except ImportError:
            logger.info("markdown not installed, PDF generation disabled")

    @property
    def is_available(self) -> bool:
        return self._available

    def convert(self, title: str, markdown_content: str) -> bytes:
        """Convert markdown to PDF bytes."""
        if not self._available:
            raise RuntimeError("PDF generation not available (install weasyprint)")

        import markdown
        from weasyprint import HTML

        # Convert markdown to HTML
        md = markdown.Markdown(extensions=["tables", "fenced_code"])
        html_content = md.convert(markdown_content)

        # Wrap in HTML document with styling
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; }}
                h2 {{ color: #555; }}
                code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background: #f4f4f4; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            {html_content}
        </body>
        </html>
        """

        # Convert to PDF
        pdf_bytes = HTML(string=html).write_pdf()
        return pdf_bytes


# =============================================================================
# Document Generator Tool Executor
# =============================================================================


class DocumentGeneratorExecutor(ToolExecutor):
    """Executor for document generation tool."""

    def __init__(self):
        self.docx_converter = MarkdownToDocxConverter()
        self.pdf_converter = MarkdownToPdfConverter()

    async def execute(self, request: ToolCallRequest) -> ToolCallResult:
        """Execute document generation."""
        start_time = time.time()

        title = request.arguments.get("title", "Document")
        content = request.arguments.get("content", "")
        format = request.arguments.get("format", "docx")

        if not content:
            return ToolCallResult(
                call_id=request.call_id,
                tool_name=request.tool_name,
                success=False,
                error="Content is required",
            )

        try:
            if format == "docx":
                if not self.docx_converter.is_available:
                    return ToolCallResult(
                        call_id=request.call_id,
                        tool_name=request.tool_name,
                        success=False,
                        error="Word document generation not available (install python-docx)",
                    )

                doc_bytes = self.docx_converter.convert(title, content)
                mime_type = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                filename = f"{self._sanitize_filename(title)}.docx"

            elif format == "pdf":
                if not self.pdf_converter.is_available:
                    return ToolCallResult(
                        call_id=request.call_id,
                        tool_name=request.tool_name,
                        success=False,
                        error="PDF generation not available (install weasyprint)",
                    )

                doc_bytes = self.pdf_converter.convert(title, content)
                mime_type = "application/pdf"
                filename = f"{self._sanitize_filename(title)}.pdf"

            elif format == "md":
                # Plain markdown output
                full_content = f"# {title}\n\n{content}"
                doc_bytes = full_content.encode("utf-8")
                mime_type = "text/markdown"
                filename = f"{self._sanitize_filename(title)}.md"

            else:
                return ToolCallResult(
                    call_id=request.call_id,
                    tool_name=request.tool_name,
                    success=False,
                    error=f"Unsupported format: {format}",
                )

            duration_ms = (time.time() - start_time) * 1000

            return ToolCallResult(
                call_id=request.call_id,
                tool_name=request.tool_name,
                success=True,
                result=f"Document '{filename}' generated successfully. The user can download it directly from the chat interface - do NOT mention any file paths or workspace directories.",
                output_files=[
                    {
                        "filename": filename,
                        "content_base64": base64.b64encode(doc_bytes).decode("utf-8"),
                        "mime_type": mime_type,
                        "size_bytes": len(doc_bytes),
                    }
                ],
                metadata={
                    "format": format,
                    "title": title,
                    "filename": filename,
                    "size_bytes": len(doc_bytes),
                    "duration_ms": duration_ms,
                },
                duration_ms=duration_ms,
            )

        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            return ToolCallResult(
                call_id=request.call_id,
                tool_name=request.tool_name,
                success=False,
                error=str(e),
                duration_ms=(time.time() - start_time) * 1000,
            )

    def _sanitize_filename(self, title: str) -> str:
        """Sanitize title for use as filename."""
        # Remove or replace invalid characters
        safe = re.sub(r'[<>:"/\\|?*]', "_", title)
        # Limit length
        return safe[:100]


# =============================================================================
# Registration Helper
# =============================================================================


def register_document_generation_tool() -> bool:
    """Register document generation tool with the global registry."""
    executor = DocumentGeneratorExecutor()

    if not executor.docx_converter.is_available:
        logger.warning("python-docx not installed, document generation tool not registered")
        return False

    register_tool(DOCUMENT_GENERATION_DEFINITION, executor)
    logger.info("Registered document generation tool")
    return True
