"""Document text extraction for various file formats (PDF, DOCX, DOC, HTML, etc.)."""
from __future__ import annotations

import asyncio
import codecs
import re
import struct
import zipfile
from collections.abc import Iterable
from html.parser import HTMLParser
from io import BytesIO
from typing import Any
from xml.parsers import expat

from ...core.exceptions import ValidationFailedError
from ...core.observability.logging import get_logger
from .ingestion_service import (
    MAX_EXTRACTED_TEXT_BYTES,
    MAX_EXTRACTED_TEXT_CHARS,
    _require_extracted_text_budget,
    _require_extracted_text_counts_budget,
)
from .utils import normalize_text

logger = get_logger(__name__)

# VLM rate-limiting state (shared across instances)
_global_vlm_semaphore: asyncio.Semaphore | None = None
_global_vlm_max_concurrent: int = 10
_global_vlm_lock = asyncio.Lock()

# OOXML is a ZIP container. These limits are checked from the central directory
# before python-docx is allowed to materialize any XML parts. The uncompressed
# byte limits intentionally share the ingestion text byte budget.
MAX_DOCX_ZIP_ENTRIES = 4_096
MAX_DOCX_ZIP_SINGLE_UNCOMPRESSED_BYTES = 16 * 1024 * 1024
MAX_DOCX_ZIP_TOTAL_UNCOMPRESSED_BYTES = MAX_EXTRACTED_TEXT_BYTES
MAX_DOCX_ZIP_COMPRESSION_RATIO = 100.0
_DECODE_CHUNK_BYTES = 64 * 1024
MAX_EXTRACTED_TEXT_PARTS = 65_536

# Format-specific materialization budgets are intentionally below the final
# text budget. They protect parsers whose APIs otherwise build a whole DOM,
# page, or table before returning control to us.
MAX_TEXT_SOURCE_BYTES = 16 * 1024 * 1024
MAX_HTML_SOURCE_BYTES = 4 * 1024 * 1024
MAX_PDF_SOURCE_BYTES = 8 * 1024 * 1024
MAX_DOCX_SOURCE_BYTES = 16 * 1024 * 1024

# pypdf parses decoded content streams into an operator-object list, so this
# cap is deliberately much smaller than the final text budget.
MAX_PDF_DECOMPRESSED_STREAM_BYTES = 2 * 1024 * 1024
MAX_PDF_PAGE_MARKERS = 500
MAX_PDF_OBJECT_MARKERS = 100_000
MAX_PDF_PAGE_TEXT_CHARS = 256_000
MAX_PDF_LINE_CHARS = 32_000
MAX_PDF_TABLES = 10_000
MAX_PDF_TABLE_ROWS = 100_000
MAX_PDF_TABLE_CELLS = 500_000

MAX_DOCX_XML_ELEMENTS = 100_000
MAX_DOCX_XML_DEPTH = 128
MAX_DOCX_XML_ATTRIBUTES = 200_000
MAX_DOCX_PARAGRAPHS = 25_000
MAX_DOCX_TABLES = 2_000
MAX_DOCX_TABLE_ROWS = 25_000
MAX_DOCX_TABLE_CELLS = 50_000

MAX_HTML_TAGS = 50_000
MAX_HTML_DEPTH = 256
MAX_HTML_ATTRIBUTES = 200_000
MAX_HTML_ATTRIBUTE_CHARS = 4_000_000

_HTML_VOID_ELEMENTS = frozenset(
    {
        "area",
        "base",
        "br",
        "col",
        "embed",
        "hr",
        "img",
        "input",
        "link",
        "meta",
        "param",
        "source",
        "track",
        "wbr",
    }
)
_POSTGRES_UNSAFE_CONTROLS = str.maketrans(
    {codepoint: None for codepoint in range(32) if codepoint not in {9, 10, 13}}
)
_ZIP_EOCD_SIGNATURE = b"PK\x05\x06"
_ZIP_EOCD_STRUCT = struct.Struct("<4s4H2LH")
_ZIP_CENTRAL_SIGNATURE = b"PK\x01\x02"
_ZIP_CENTRAL_STRUCT = struct.Struct("<4s6H3L5H2L")
_ZIP_MAX_COMMENT_BYTES = 65_535
_ZIP64_EXTRA_FIELD_ID = 0x0001


def _bounded_utf8_size(value: str, *, limit: int) -> int:
    total = 0
    for offset in range(0, len(value), _DECODE_CHUNK_BYTES):
        total += len(value[offset : offset + _DECODE_CHUNK_BYTES].encode("utf-8"))
        if total > limit:
            break
    return total


def _require_source_bytes(content: bytes, *, limit: int, format_name: str) -> None:
    if len(content) > limit:
        raise ValidationFailedError(
            f"{format_name} source exceeds the {limit} byte parser limit"
        )


def _preflight_docx_central_directory(content: bytes) -> None:
    """Bound ZIP entry materialization before ``zipfile.ZipFile`` is created."""
    search_start = max(
        len(content) - _ZIP_EOCD_STRUCT.size - _ZIP_MAX_COMMENT_BYTES,
        0,
    )
    eocd_offset = content.rfind(_ZIP_EOCD_SIGNATURE, search_start)
    if eocd_offset < 0 or eocd_offset + _ZIP_EOCD_STRUCT.size > len(content):
        raise ValidationFailedError("DOCX archive is missing a valid central directory")
    (
        _signature,
        disk_number,
        central_directory_disk,
        entries_on_disk,
        total_entries,
        central_directory_size,
        central_directory_offset,
        comment_size,
    ) = _ZIP_EOCD_STRUCT.unpack_from(content, eocd_offset)
    if eocd_offset + _ZIP_EOCD_STRUCT.size + comment_size != len(content):
        raise ValidationFailedError("DOCX archive has an invalid central directory trailer")
    if disk_number or central_directory_disk or entries_on_disk != total_entries:
        raise ValidationFailedError("Multi-disk DOCX archives are not supported")
    if (
        total_entries == 0xFFFF
        or central_directory_size == 0xFFFFFFFF
        or central_directory_offset == 0xFFFFFFFF
    ):
        raise ValidationFailedError("ZIP64 DOCX archives are not supported")
    if (
        central_directory_size > MAX_DOCX_SOURCE_BYTES
        or central_directory_offset + central_directory_size != eocd_offset
    ):
        raise ValidationFailedError("DOCX archive central directory is invalid or oversized")

    directory_end = central_directory_offset + central_directory_size
    cursor = central_directory_offset
    actual_entries = 0
    while cursor < directory_end:
        if cursor + _ZIP_CENTRAL_STRUCT.size > directory_end:
            raise ValidationFailedError("DOCX central directory record is truncated")
        (
            signature,
            _version_made_by,
            _version_needed,
            _flags,
            _compression,
            _modified_time,
            _modified_date,
            _crc32,
            compressed_size,
            uncompressed_size,
            filename_size,
            extra_size,
            record_comment_size,
            disk_start,
            _internal_attributes,
            _external_attributes,
            local_header_offset,
        ) = _ZIP_CENTRAL_STRUCT.unpack_from(content, cursor)
        if signature != _ZIP_CENTRAL_SIGNATURE:
            raise ValidationFailedError("DOCX central directory record signature is invalid")
        if disk_start:
            raise ValidationFailedError("Multi-disk DOCX entries are not supported")
        if (
            compressed_size == 0xFFFFFFFF
            or uncompressed_size == 0xFFFFFFFF
            or local_header_offset == 0xFFFFFFFF
        ):
            raise ValidationFailedError("ZIP64 DOCX entries are not supported")
        record_end = (
            cursor
            + _ZIP_CENTRAL_STRUCT.size
            + filename_size
            + extra_size
            + record_comment_size
        )
        if record_end > directory_end:
            raise ValidationFailedError("DOCX central directory record exceeds its bounds")
        if filename_size == 0 or local_header_offset >= central_directory_offset:
            raise ValidationFailedError("DOCX central directory record is invalid")

        extra_cursor = cursor + _ZIP_CENTRAL_STRUCT.size + filename_size
        extra_end = extra_cursor + extra_size
        while extra_cursor < extra_end:
            if extra_cursor + 4 > extra_end:
                raise ValidationFailedError("DOCX central directory extra field is malformed")
            field_id, field_size = struct.unpack_from("<HH", content, extra_cursor)
            extra_cursor += 4
            if extra_cursor + field_size > extra_end:
                raise ValidationFailedError("DOCX central directory extra field is malformed")
            if field_id == _ZIP64_EXTRA_FIELD_ID:
                raise ValidationFailedError("ZIP64 DOCX entries are not supported")
            extra_cursor += field_size

        actual_entries += 1
        if actual_entries > MAX_DOCX_ZIP_ENTRIES:
            raise ValidationFailedError(
                f"DOCX archive exceeds the {MAX_DOCX_ZIP_ENTRIES} entry limit"
            )
        cursor = record_end

    if cursor != directory_end or actual_entries != total_entries:
        raise ValidationFailedError("DOCX central directory entry count is inconsistent")


def _iter_lines(value: str) -> Iterable[str]:
    """Yield newline-delimited spans without allocating a full split list."""
    start = 0
    while True:
        end = value.find("\n", start)
        if end < 0:
            yield value[start:]
            return
        yield value[start:end]
        start = end + 1


class _HTMLStructureGuard(HTMLParser):
    """Streaming HTML complexity preflight used before BeautifulSoup DOM build."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tag_count = 0
        self.attribute_count = 0
        self.attribute_chars = 0
        self._open_tags: list[str] = []

    def _charge_tag(self, attrs: list[tuple[str, str | None]]) -> None:
        self.tag_count += 1
        self.attribute_count += len(attrs)
        self.attribute_chars += sum(
            len(name or "") + len(value or "") for name, value in attrs
        )
        if self.tag_count > MAX_HTML_TAGS:
            raise ValidationFailedError(
                f"HTML exceeds the {MAX_HTML_TAGS} tag parser limit"
            )
        if self.attribute_count > MAX_HTML_ATTRIBUTES:
            raise ValidationFailedError(
                f"HTML exceeds the {MAX_HTML_ATTRIBUTES} attribute parser limit"
            )
        if self.attribute_chars > MAX_HTML_ATTRIBUTE_CHARS:
            raise ValidationFailedError(
                "HTML attributes exceed the parser text budget"
            )

    def handle_starttag(
        self,
        tag: str,
        attrs: list[tuple[str, str | None]],
    ) -> None:
        normalized = tag.lower()
        self._charge_tag(attrs)
        if normalized not in _HTML_VOID_ELEMENTS:
            self._open_tags.append(normalized)
            if len(self._open_tags) > MAX_HTML_DEPTH:
                raise ValidationFailedError(
                    f"HTML exceeds the {MAX_HTML_DEPTH} nesting depth limit"
                )

    def handle_startendtag(
        self,
        _tag: str,
        attrs: list[tuple[str, str | None]],
    ) -> None:
        self._charge_tag(attrs)

    def handle_endtag(self, tag: str) -> None:
        normalized = tag.lower()
        for index in range(len(self._open_tags) - 1, -1, -1):
            if self._open_tags[index] == normalized:
                del self._open_tags[index:]
                break


def _preflight_html_structure(html: str) -> None:
    if _bounded_utf8_size(html, limit=MAX_HTML_SOURCE_BYTES) > MAX_HTML_SOURCE_BYTES:
        raise ValidationFailedError(
            f"HTML source exceeds the {MAX_HTML_SOURCE_BYTES} byte parser limit"
        )
    guard = _HTMLStructureGuard()
    try:
        for offset in range(0, len(html), _DECODE_CHUNK_BYTES):
            guard.feed(html[offset : offset + _DECODE_CHUNK_BYTES])
        guard.close()
    except ValidationFailedError:
        raise
    except Exception as exc:
        raise ValidationFailedError(f"HTML structure preflight failed: {exc}") from exc


def _preflight_pdf_bytes(content: bytes) -> None:
    _require_source_bytes(
        content,
        limit=MAX_PDF_SOURCE_BYTES,
        format_name="PDF",
    )
    for page_markers, _match in enumerate(
        re.finditer(rb"/Type\s*/Page\b", content),
        start=1,
    ):
        if page_markers > MAX_PDF_PAGE_MARKERS:
            raise ValidationFailedError(
                f"PDF exceeds the {MAX_PDF_PAGE_MARKERS} page marker limit"
            )
    for object_markers, _match in enumerate(
        re.finditer(rb"\b\d+\s+\d+\s+obj\b", content),
        start=1,
    ):
        if object_markers > MAX_PDF_OBJECT_MARKERS:
            raise ValidationFailedError(
                f"PDF exceeds the {MAX_PDF_OBJECT_MARKERS} object marker limit"
            )


class _PDFTableBudget:
    def __init__(self) -> None:
        self.tables = 0
        self.rows = 0
        self.cells = 0

    def begin_table(self) -> None:
        self.tables += 1
        if self.tables > MAX_PDF_TABLES:
            raise ValidationFailedError(
                f"PDF exceeds the {MAX_PDF_TABLES} table limit"
            )

    def charge_row(self, cells: int) -> None:
        self.rows += 1
        self.cells += cells
        if self.rows > MAX_PDF_TABLE_ROWS:
            raise ValidationFailedError(
                f"PDF exceeds the {MAX_PDF_TABLE_ROWS} table row limit"
            )
        if self.cells > MAX_PDF_TABLE_CELLS:
            raise ValidationFailedError(
                f"PDF exceeds the {MAX_PDF_TABLE_CELLS} table cell limit"
            )


class _DOCXXMLBudget:
    """SAX-style OOXML budget shared across all XML package parts."""

    def __init__(self) -> None:
        self.elements = 0
        self.depth = 0
        self.attributes = 0
        self.text_chars = 0
        self.text_bytes = 0
        self.structures = {"p": 0, "tbl": 0, "tr": 0, "tc": 0}

    def start_element(self, name: str, attrs: dict[str, str]) -> None:
        self.elements += 1
        self.depth += 1
        self.attributes += len(attrs)
        if self.elements > MAX_DOCX_XML_ELEMENTS:
            raise ValidationFailedError(
                f"DOCX XML exceeds the {MAX_DOCX_XML_ELEMENTS} element limit"
            )
        if self.depth > MAX_DOCX_XML_DEPTH:
            raise ValidationFailedError(
                f"DOCX XML exceeds the {MAX_DOCX_XML_DEPTH} depth limit"
            )
        if self.attributes > MAX_DOCX_XML_ATTRIBUTES:
            raise ValidationFailedError(
                f"DOCX XML exceeds the {MAX_DOCX_XML_ATTRIBUTES} attribute limit"
            )
        local_name = name.rsplit("}", 1)[-1]
        if local_name in self.structures:
            self.structures[local_name] += 1
            limits = {
                "p": MAX_DOCX_PARAGRAPHS,
                "tbl": MAX_DOCX_TABLES,
                "tr": MAX_DOCX_TABLE_ROWS,
                "tc": MAX_DOCX_TABLE_CELLS,
            }
            if self.structures[local_name] > limits[local_name]:
                raise ValidationFailedError(
                    f"DOCX XML exceeds the {limits[local_name]} {local_name} limit"
                )

    def end_element(self, _name: str) -> None:
        self.depth = max(self.depth - 1, 0)

    def character_data(self, value: str) -> None:
        self.text_chars += len(value)
        remaining = MAX_EXTRACTED_TEXT_BYTES - self.text_bytes
        self.text_bytes += _bounded_utf8_size(value, limit=remaining)
        _require_extracted_text_counts_budget(
            self.text_chars,
            self.text_bytes,
        )


def _validate_docx_xml_stream(source: Any, budget: _DOCXXMLBudget) -> None:
    parser = expat.ParserCreate(namespace_separator="}")
    parser.StartElementHandler = budget.start_element
    parser.EndElementHandler = budget.end_element
    parser.CharacterDataHandler = budget.character_data

    def reject_doctype(*_args: Any) -> None:
        raise ValidationFailedError("DOCX XML document types are not supported")

    parser.StartDoctypeDeclHandler = reject_doctype
    parser.ExternalEntityRefHandler = lambda *_args: 0
    try:
        while True:
            chunk = source.read(_DECODE_CHUNK_BYTES)
            if not chunk:
                break
            parser.Parse(chunk, False)
        parser.Parse(b"", True)
    except ValidationFailedError:
        raise
    except (expat.ExpatError, OSError, ValueError) as exc:
        raise ValidationFailedError(f"DOCX contains invalid XML: {exc}") from exc


class _BoundedTextBuilder:
    """Incrementally assemble extracted text under the ingestion budget."""

    def __init__(
        self,
        *,
        max_chars: int | None = None,
        max_bytes: int | None = None,
        max_parts: int = MAX_EXTRACTED_TEXT_PARTS,
    ) -> None:
        self._max_chars = MAX_EXTRACTED_TEXT_CHARS if max_chars is None else max_chars
        self._max_bytes = MAX_EXTRACTED_TEXT_BYTES if max_bytes is None else max_bytes
        self._max_parts = max_parts
        self._parts: list[str] = []
        self._chars = 0
        self._bytes = 0

    def append(self, value: Any, *, separator: str = "") -> None:
        text = str(value or "")
        if not text:
            return
        prefix = separator if self._parts and separator else ""
        self._charge(prefix)
        self._charge(text)
        fragment = prefix + text
        if len(self._parts) >= self._max_parts:
            raise ValidationFailedError(
                f"extracted text exceeds the {self._max_parts} fragment limit"
            )
        self._parts.append(fragment)

    def _charge(self, text: str) -> None:
        self._chars += len(text)
        if self._chars > self._max_chars:
            raise ValidationFailedError(
                f"extracted text exceeds the {self._max_chars} character limit"
            )
        remaining_bytes = self._max_bytes - self._bytes
        self._bytes += _bounded_utf8_size(text, limit=remaining_bytes)
        if self._bytes > self._max_bytes:
            raise ValidationFailedError(
                f"extracted text exceeds the {self._max_bytes} byte limit"
            )

    def build(self) -> str:
        value = "".join(self._parts)
        # Recheck the actual materialized result with the canonical ingestion
        # helper so incremental and downstream enforcement cannot drift.
        return _require_extracted_text_budget(value)


def _append_markdown_block(
    builder: _BoundedTextBuilder,
    lines: Iterable[str],
    *,
    block_separator: str,
) -> None:
    first = True
    for line in lines:
        if not line:
            continue
        builder.append(line, separator=block_separator if first else "\n")
        first = False


class DocumentProcessor:
    """Extracts and normalises text from various document formats."""

    def __init__(self, settings: Any, vlm_service: Any | None = None):
        self.settings = settings
        self.vlm_service = vlm_service

    # ------------------------------------------------------------------
    # VLM callback
    # ------------------------------------------------------------------

    def create_vlm_callback(self):
        """Create a VLM callback for document processing.

        Uses a global semaphore to limit concurrent VLM API calls across all
        document processing tasks. This prevents overwhelming the API when
        multiple documents are processed simultaneously.
        """
        global _global_vlm_semaphore, _global_vlm_max_concurrent

        vlm = self.vlm_service
        if vlm is None:
            return None

        # Initialize global semaphore if not done yet
        vlm_max_concurrent = self.settings.knowledge.vlm_max_concurrent
        if _global_vlm_semaphore is None or _global_vlm_max_concurrent != vlm_max_concurrent:
            _global_vlm_max_concurrent = vlm_max_concurrent
            _global_vlm_semaphore = asyncio.Semaphore(vlm_max_concurrent)
            logger.info(
                f"Initialized global VLM semaphore with max_concurrent={vlm_max_concurrent}"
            )

        semaphore = _global_vlm_semaphore

        async def _vlm_extract_text(image_bytes: bytes, lang: str) -> str:
            prompt = (
                "Extract ALL text from this document page exactly as written. "
                "Preserve the original structure, paragraphs, and formatting. "
                "Do not summarize or interpret — output only the raw text content."
            )
            if lang == "ar":
                prompt = (
                    "استخرج جميع النصوص من هذه الصفحة كما هي مكتوبة بالضبط. "
                    "حافظ على الهيكل الأصلي والفقرات والتنسيق. "
                    "لا تلخص أو تفسر — أخرج فقط المحتوى النصي الخام."
                )
            try:
                # Use global semaphore to limit concurrent VLM calls
                async with semaphore:
                    result = await vlm.describe_image(
                        image_bytes=image_bytes,
                        prompt=prompt,
                        image_type="document",
                        max_tokens=2000,
                    )
                    return result.description
            except Exception as e:
                logger.warning(f"VLM text extraction failed: {e}")
                return ""

        return _vlm_extract_text

    # ------------------------------------------------------------------
    # Sanitisation helpers
    # ------------------------------------------------------------------

    def _sanitize_text_for_db(self, text: str) -> str:
        """Remove NULL bytes and other characters that PostgreSQL cannot handle."""
        if not text:
            return ""
        _require_extracted_text_budget(text)
        # The old per-character list could allocate millions of Python objects.
        # ``translate`` preserves the exact rule: remove ASCII controls except
        # tab/newline/carriage-return, and retain all code points above 31.
        return _require_extracted_text_budget(text.translate(_POSTGRES_UNSAFE_CONTROLS))

    def decode_text_bytes(self, content: bytes) -> str:
        _require_source_bytes(
            content,
            limit=MAX_TEXT_SOURCE_BYTES,
            format_name="text",
        )
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
            try:
                decoder = codecs.getincrementaldecoder(enc)(errors="strict")
                builder = _BoundedTextBuilder()
                for offset in range(0, len(content), _DECODE_CHUNK_BYTES):
                    decoded = decoder.decode(
                        content[offset : offset + _DECODE_CHUNK_BYTES],
                        final=False,
                    )
                    builder.append(decoded)
                builder.append(decoder.decode(b"", final=True))
                # Sanitize for PostgreSQL compatibility
                return self._sanitize_text_for_db(builder.build())
            except UnicodeDecodeError:
                continue
        raise ValidationFailedError("Unable to decode uploaded file as text")

    # ------------------------------------------------------------------
    # PDF extraction
    # ------------------------------------------------------------------

    def clean_pdf_content(self, text: str) -> str:
        """Clean PDF extracted content, removing TOC lines and noise."""
        if not text:
            return ""
        _require_extracted_text_budget(text)

        builder = _BoundedTextBuilder()

        # Track if we're in TOC section
        toc_indicators = 0

        for line in _iter_lines(text):
            if len(line) > MAX_PDF_LINE_CHARS:
                raise ValidationFailedError(
                    f"PDF line exceeds the {MAX_PDF_LINE_CHARS} character limit"
                )
            line = line.strip()
            if not line:
                continue

            # Skip lines that look like TOC entries - VERY aggressive patterns
            suffix_end = len(line.rstrip())
            suffix_start = suffix_end
            while suffix_start > 0 and line[suffix_start - 1].isdigit():
                suffix_start -= 1
            prefix = line[:suffix_start].rstrip()
            has_page_suffix = suffix_start < suffix_end
            if has_page_suffix and (
                ".." in prefix
                or "··" in prefix
                or "…" in prefix
                or prefix.count(".") >= 2
            ):
                toc_indicators += 1
                continue

            # Skip lines starting with dots (like "......2")
            leading_remainder = line.lstrip(".·… ")
            if leading_remainder != line and leading_remainder[:1].isdigit():
                toc_indicators += 1
                continue

            # Skip lines that contain excessive dots anywhere
            dot_count = line.count(".") + line.count("·") + line.count("…")
            if len(line) > 5 and dot_count > 3 and dot_count / len(line) > 0.15:
                # More than 3 dots in a short line, likely TOC
                toc_indicators += 1
                continue

            # Skip very short lines that are just page numbers or section numbers
            if re.match(r"^[\d\.\s]+$", line) and len(line) < 10:
                continue

            # Skip lines that look like "2.1 2.1标题..."
            if re.match(r"^\d+(\.\d+)*\s+\d+(\.\d+)*", line):
                continue

            # Clean up remaining dots sequences in the line
            line = re.sub(r"\.{3,}", " ", line)
            line = re.sub(r"(\.\s+){2,}", " ", line)
            line = re.sub(r"·{2,}", " ", line)
            line = re.sub(r"…{1,}", " ", line)

            # Clean up repeated spaces
            line = re.sub(r"\s{2,}", " ", line)

            line = line.strip()
            if line and len(line) > 2:  # Skip very short remnants
                builder.append(line, separator="\n")

        result = builder.build()

        # If a large portion of content was TOC-like, we may have a TOC-heavy doc
        # Log this for debugging
        if toc_indicators > 10:
            logger.info(f"Cleaned {toc_indicators} TOC-like lines from PDF")

        return _require_extracted_text_budget(result)

    def extract_text_from_pdf_bytes(self, content: bytes) -> str:
        """Extract bounded text page-by-page with guarded pypdf decompression."""
        _preflight_pdf_bytes(content)
        try:
            import pypdf  # type: ignore
        except ImportError as exc:
            raise ValidationFailedError(
                "PDF parsing requires pypdf (pip install pypdf)"
            ) from exc

        for limit_name in (
            "ZLIB_MAX_OUTPUT_LENGTH",
            "LZW_MAX_OUTPUT_LENGTH",
            "RUN_LENGTH_MAX_OUTPUT_LENGTH",
            "JBIG2_MAX_OUTPUT_LENGTH",
            "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
        ):
            current_limit = int(getattr(pypdf.filters, limit_name, 0) or 0)
            if current_limit <= 0 or current_limit > MAX_PDF_DECOMPRESSED_STREAM_BYTES:
                setattr(
                    pypdf.filters,
                    limit_name,
                    MAX_PDF_DECOMPRESSED_STREAM_BYTES,
                )

        try:
            reader = pypdf.PdfReader(BytesIO(content))
            if len(reader.pages) > MAX_PDF_PAGE_MARKERS:
                raise ValidationFailedError(
                    f"PDF exceeds the {MAX_PDF_PAGE_MARKERS} page limit"
                )
            builder = _BoundedTextBuilder()
            for page in reader.pages:
                text = page.extract_text() or ""
                if len(text) > MAX_PDF_PAGE_TEXT_CHARS:
                    raise ValidationFailedError(
                        "PDF page text exceeds the parser output limit"
                    )
                builder.append(text, separator="\n")
            text = builder.build()
        except ValidationFailedError:
            raise
        except Exception as exc:
            raise ValidationFailedError(f"Failed to parse PDF: {exc}") from exc

        if not text or not text.strip():
            raise ValidationFailedError("Failed to extract any text from PDF")

        text = self._sanitize_text_for_db(normalize_text(text))
        return _require_extracted_text_budget(self.clean_pdf_content(text))

    def ocr_pdf_bytes(self, content: bytes) -> str:
        """Fail closed: the legacy OCR API materializes an unbounded whole result."""
        del content
        raise ValidationFailedError(
            "OCR text extraction is disabled in the bounded text-only pipeline"
        )

    def extract_pdf_with_pdfplumber(self, pdf_stream) -> str:
        """Fail closed: pdfplumber materializes page geometry and all tables."""
        del pdf_stream
        raise ValidationFailedError(
            "pdfplumber table extraction is disabled in the bounded text-only pipeline"
        )

    def _iter_pdf_table_markdown_lines(
        self,
        table: Iterable[Iterable[Any]],
        *,
        budget: _PDFTableBudget,
    ) -> Iterable[str]:
        """Yield a PDF table row-by-row so extracted text stays budgeted."""
        emitted_rows = 0
        for row in table:
            cleaned = [
                str(cell or "").strip().replace("|", "\\|").replace("\n", " ")
                for cell in row
            ]
            if not any(cleaned):
                continue
            budget.charge_row(len(cleaned))
            yield "| " + " | ".join(cleaned) + " |"
            if emitted_rows == 0:
                yield "| " + " | ".join(["---"] * len(cleaned)) + " |"
            emitted_rows += 1

    def pdf_table_to_markdown(self, table: list[list]) -> str:
        """Convert a PDF table (list of rows) to Markdown format."""
        builder = _BoundedTextBuilder()
        budget = _PDFTableBudget()
        budget.begin_table()
        _append_markdown_block(
            builder,
            self._iter_pdf_table_markdown_lines(table, budget=budget),
            block_separator="",
        )
        return builder.build()

    # ------------------------------------------------------------------
    # DOCX / DOC extraction
    # ------------------------------------------------------------------

    def _validate_docx_archive(self, content: bytes) -> None:
        """Reject oversized or highly compressed OOXML before python-docx loads it."""
        _require_source_bytes(
            content,
            limit=MAX_DOCX_SOURCE_BYTES,
            format_name="DOCX",
        )
        _preflight_docx_central_directory(content)
        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_DOCX_ZIP_ENTRIES:
                    raise ValidationFailedError(
                        f"DOCX archive exceeds the {MAX_DOCX_ZIP_ENTRIES} entry limit"
                    )
                if len({entry.filename for entry in entries}) != len(entries):
                    raise ValidationFailedError("DOCX archive contains duplicate entries")

                total_uncompressed = 0
                for entry in entries:
                    if entry.flag_bits & 0x1:
                        raise ValidationFailedError(
                            "Encrypted DOCX archive entries are not supported"
                        )
                    if entry.file_size < 0 or entry.compress_size < 0:
                        raise ValidationFailedError(
                            "DOCX archive contains an invalid entry size"
                        )
                    if entry.file_size > MAX_DOCX_ZIP_SINGLE_UNCOMPRESSED_BYTES:
                        raise ValidationFailedError(
                            "DOCX archive entry exceeds the uncompressed byte limit"
                        )
                    total_uncompressed += entry.file_size
                    if total_uncompressed > MAX_DOCX_ZIP_TOTAL_UNCOMPRESSED_BYTES:
                        raise ValidationFailedError(
                            "DOCX archive exceeds the total uncompressed byte limit"
                        )
                    if entry.file_size:
                        if entry.compress_size == 0:
                            raise ValidationFailedError(
                                "DOCX archive entry has an invalid compression ratio"
                            )
                        ratio = entry.file_size / entry.compress_size
                        if ratio > MAX_DOCX_ZIP_COMPRESSION_RATIO:
                            raise ValidationFailedError(
                                "DOCX archive entry exceeds the compression ratio limit"
                            )

                xml_budget = _DOCXXMLBudget()
                for entry in entries:
                    lowered = entry.filename.lower()
                    if not lowered.endswith((".xml", ".rels")):
                        continue
                    with archive.open(entry, "r") as source:
                        _validate_docx_xml_stream(source, xml_budget)
        except ValidationFailedError:
            raise
        except (zipfile.BadZipFile, OSError, ValueError) as exc:
            raise ValidationFailedError("DOCX is not a valid OOXML ZIP archive") from exc

    def extract_text_from_docx_bytes(self, content: bytes) -> str:
        """Extract text from DOCX with table-to-markdown conversion."""
        self._validate_docx_archive(content)
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise ValidationFailedError(
                "DOCX parsing requires python-docx (pip install python-docx)"
            ) from exc

        try:
            doc = Document(BytesIO(content))
            builder = _BoundedTextBuilder()
            iter_content = getattr(doc, "iter_inner_content", None)
            if not callable(iter_content):
                raise ValidationFailedError(
                    "bounded DOCX parsing requires python-docx.iter_inner_content"
                )
            for block in iter_content():
                if hasattr(block, "rows"):
                    _append_markdown_block(
                        builder,
                        self._iter_docx_table_markdown_lines(block),
                        block_separator="\n\n",
                    )
                    continue
                text = str(getattr(block, "text", "") or "").strip()
                if text:
                    builder.append(text, separator="\n")

            text = normalize_text(builder.build())
            if not text:
                raise ValidationFailedError("DOCX parsed but no text extracted")
            return _require_extracted_text_budget(self._sanitize_text_for_db(text))
        except ValidationFailedError:
            raise
        except Exception as exc:
            raise ValidationFailedError(f"Failed to parse DOCX: {exc}") from exc

    def _iter_docx_table_markdown_lines(self, table: Any) -> Iterable[str]:
        """Yield a python-docx table incrementally."""
        emitted_rows = 0
        for row in getattr(table, "rows", ()) or ():
            cells = [
                str(getattr(cell, "text", "") or "")
                .strip()
                .replace("|", "\\|")
                .replace("\n", " ")
                for cell in (getattr(row, "cells", ()) or ())
            ]
            if not cells:
                continue
            yield "| " + " | ".join(cells) + " |"
            if emitted_rows == 0:
                yield "| " + " | ".join(["---"] * len(cells)) + " |"
            emitted_rows += 1

    def table_to_markdown(self, table) -> str:
        """Convert a python-docx table to Markdown format."""
        try:
            builder = _BoundedTextBuilder()
            _append_markdown_block(
                builder,
                self._iter_docx_table_markdown_lines(table),
                block_separator="",
            )
            return builder.build()
        except ValidationFailedError:
            raise
        except Exception:
            return ""

    def parse_table_row(self, row, total_cols: int) -> list[str]:
        """Parse a table row into a list of cell texts."""
        cells = list(getattr(row, "cells", []) or [])
        row_cells = [""] * total_cols
        col_idx = 0

        for cell in cells:
            if col_idx >= total_cols:
                break
            # Skip already filled cells (from previous merged cells)
            while col_idx < total_cols and row_cells[col_idx]:
                col_idx += 1
            if col_idx >= total_cols:
                break

            # Get cell text
            cell_text = str(getattr(cell, "text", "") or "").strip()
            # Clean up cell text for markdown (escape pipes, remove newlines)
            cell_text = cell_text.replace("|", "\\|").replace("\n", " ")

            # Handle grid span (column merging)
            grid_span = getattr(cell, "grid_span", 1) or 1
            for i in range(grid_span):
                if col_idx + i < total_cols:
                    row_cells[col_idx + i] = cell_text if i == 0 else ""
            col_idx += grid_span

        return row_cells

    def extract_text_from_doc_bytes(self, content: bytes) -> str:
        """Fail closed: textract exposes only an unbounded whole-output API."""
        del content
        raise ValidationFailedError(
            "Legacy DOC extraction is disabled; convert the document to DOCX or text"
        )

    # ------------------------------------------------------------------
    # HTML extraction
    # ------------------------------------------------------------------

    def extract_text_from_html(self, html: str) -> str:
        """Extract text from HTML with improved handling of various content types."""
        html = _require_extracted_text_budget(html or "")
        _preflight_html_structure(html)
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except Exception as exc:
            raise ValidationFailedError(
                "HTML parsing requires beautifulsoup4 (pip install beautifulsoup4 lxml)"
            ) from exc

        soup = BeautifulSoup(html or "", "lxml")

        # Try to find main content area
        main_content = None
        for selector in [
            "main",
            "article",
            "[role='main']",
            ".content",
            "#content",
            ".post",
            ".article",
        ]:
            try:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            except Exception:
                continue

        # Use main content if found, otherwise use full body
        content_root = main_content or soup.body or soup

        builder = _BoundedTextBuilder()
        ignored_names = {
            "script",
            "style",
            "noscript",
            "header",
            "footer",
            "nav",
            "aside",
            "iframe",
            "form",
        }

        def is_ignored(element: Any) -> bool:
            current = element
            while current is not None and current is not content_root:
                if str(getattr(current, "name", "") or "").lower() in ignored_names:
                    return True
                current = getattr(current, "parent", None)
            return False

        # Extract title
        title_tag = soup.find("title")
        if title_tag and title_tag.string:
            builder.append(f"# {title_tag.string.strip()}", separator="\n\n")

        # Descendants is a generator; unlike find_all it does not first build a
        # document-sized ResultSet. Table rows are handled in the same pass.
        content_names = {
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "p",
            "li",
            "blockquote",
            "pre",
            "code",
        }
        for element in content_root.descendants:
            tag_name = str(getattr(element, "name", "") or "").lower()
            if not tag_name or is_ignored(element):
                continue
            if tag_name == "tr":
                cells: list[str] = []
                for child in getattr(element, "children", ()):
                    if str(getattr(child, "name", "") or "").lower() not in {"th", "td"}:
                        continue
                    cells.append(child.get_text(separator=" ", strip=True))
                if any(cells):
                    builder.append(
                        " | ".join(cell if cell else "-" for cell in cells),
                        separator="\n",
                    )
                continue
            if tag_name not in content_names:
                continue
            text = element.get_text(separator=" ", strip=True)
            if text:
                if tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    level = int(tag_name[1])
                    builder.append("#" * level + " " + text, separator="\n\n")
                elif tag_name == "li":
                    builder.append("• " + text, separator="\n\n")
                else:
                    builder.append(text, separator="\n\n")

        structured_text = builder.build()

        # Fallback: if no structured content found, use simple text extraction
        if not structured_text:
            fallback = _BoundedTextBuilder()
            for text in content_root.stripped_strings:
                fallback.append(text, separator="\n")
            result = self._sanitize_text_for_db(normalize_text(fallback.build()))
            return _require_extracted_text_budget(result)

        result = normalize_text(structured_text)
        return _require_extracted_text_budget(self._sanitize_text_for_db(result))

    # ------------------------------------------------------------------
    # Unified dispatcher
    # ------------------------------------------------------------------

    def extract_text_from_bytes(
        self,
        content: bytes,
        filename: str | None = None,
        mime_type: str | None = None,
    ) -> tuple[str, str]:
        name = (filename or "").strip().lower()
        mime = (mime_type or "").strip().lower()

        # Legacy Office (.doc) is OLE2 Compound Document.
        if (
            content.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            or name.endswith(".doc")
            or "application/msword" in mime
        ):
            text = self.extract_text_from_doc_bytes(content)
            return _require_extracted_text_budget(text), "application/msword"

        # PDF
        if content.startswith(b"%PDF") or name.endswith(".pdf") or "application/pdf" in mime:
            text = self.extract_text_from_pdf_bytes(content)
            return _require_extracted_text_budget(text), "application/pdf"

        # DOCX (OOXML zip)
        if (
            content.startswith(b"PK\x03\x04")
            or name.endswith(".docx")
            or "wordprocessingml.document" in mime
        ):
            text = self.extract_text_from_docx_bytes(content)
            return (
                _require_extracted_text_budget(text),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # HTML (primarily for URL ingestion)
        if "text/html" in mime or name.endswith(".html") or name.endswith(".htm"):
            _require_source_bytes(
                content,
                limit=MAX_HTML_SOURCE_BYTES,
                format_name="HTML",
            )
            decoded = self.decode_text_bytes(content)
            text = self.extract_text_from_html(decoded)
            return _require_extracted_text_budget(text), "text/html"

        # Markdown
        if name.endswith(".md") or mime in {"text/markdown", "text/x-markdown"}:
            text = self.decode_text_bytes(content)
            return _require_extracted_text_budget(text), "text/markdown"

        # Plain text fallback
        decoded = self.decode_text_bytes(content)
        return _require_extracted_text_budget(decoded), (mime_type or "text/plain")
